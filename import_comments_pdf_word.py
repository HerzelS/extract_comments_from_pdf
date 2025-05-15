"""
PDF Comment Extractor

This script extracts comments and highlighted text (as comment scope) from a PDF file
and exports them into a Word document in a landscape table format.

Dependencies:
    - pymupdf (fitz)
    - python-docx

Install with:
    pip install pymupdf python-docx
"""

import fitz  # PyMuPDF for reading PDF annotations
from docx import Document  # For creating the Word document
from docx.enum.section import WD_ORIENT  # For setting page orientation


def extract_comments_to_word(pdf_path, word_output_path):
    """
    Extracts comments from a PDF file and writes them to a Word document table.

    Each comment is added as a new row with the following columns:
    Author, Page, Paragraph, Comment Scope (highlighted text), Comment Text, Team Response.

    Parameters:
        pdf_path (str): Path to the input PDF file.
        word_output_path (str): Path to the output Word (.docx) file.

    Notes:
        - Comment Scope contains the text highlighted during annotation.
        - Paragraph is marked as "N/A" because paragraph info is not available in PDFs.
        - Only highlight annotations are currently supported for scope extraction.
    """
    
    # Open the input PDF document
    doc = fitz.open(pdf_path)

    # Create a new Word document
    word_doc = Document()
    
    # Set document orientation to landscape
    section = word_doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # Add a header with metadata
    header = section.header
    header.paragraphs[0].text = (
        f"Comments extracted from: {pdf_path}\n"
        f"Created by: Python Script\n"
        f"Creation date:"
    )
    
    # Create a table with 6 columns
    table = word_doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Set column headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Author'
    hdr_cells[1].text = 'Page'
    hdr_cells[2].text = 'Paragraph'
    hdr_cells[3].text = 'Comment Scope'
    hdr_cells[4].text = 'Comment Text'
    hdr_cells[5].text = 'Team Response'

    # Iterate through each page and extract annotations
    for page_num, page in enumerate(doc, start=1):
        for annot in page.annots() or []:
            author = annot.info.get("title", "Unknown")  # Author of the comment
            comment_text = annot.info.get("content", "")  # The comment text
            scope_text = ""  # To hold the highlighted text (if any)

            # If annotation is a highlight (type 8), extract highlighted text
            if annot.type[0] == 8:
                quads = annot.vertices
                quad_count = int(len(quads) / 4)
                words = page.get_text("words")  # Get all words on the page
                words_in_quad = []

                # Check each quad and collect intersecting words
                for i in range(quad_count):
                    rect = fitz.Quad(quads[i * 4: i * 4 + 4]).rect
                    for word in words:
                        word_rect = fitz.Rect(word[:4])
                        if rect.intersects(word_rect):
                            words_in_quad.append(word[4])

                scope_text = " ".join(words_in_quad).strip()

            # Add the comment data to a new row in the table
            row_cells = table.add_row().cells
            row_cells[0].text = author
            row_cells[1].text = str(page_num)
            row_cells[2].text = "N/A"  # Paragraph info not available in PDFs
            row_cells[3].text = scope_text if scope_text else "N/A"
            row_cells[4].text = comment_text
            row_cells[5].text = ""  # Placeholder for Team Response

    # Save the Word document
    word_doc.save(word_output_path)
    print(f"Extracted comments saved to: {word_output_path}")


# Example usage - rename the file path as needed
extract_comments_to_word("file.pdf", "comments_output.docx")
