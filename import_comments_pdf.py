import fitz # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT

def extract_comments_to_word(pdf_path, word_output_path):
    # Open the PDF
    doc = fitz.open(pdf_path)

    # Create Word Document
    word_doc = Document()

    # Set landscape orientation
    section = word_doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Add header
    header = section.header
    header.paragraphs[0].text = f"Comments extracted from: {pdf_path}\nCreated by: Python Script\nCreation date:"


    # Add a table with 6 columns
    table = word_doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Author"
    hdr_cells[1].text = "Page"
    hdr_cells[2].text = "Paragraph"
    hdr_cells[3].text = "Comment Scope"
    hdr_cells[4].text = "Comment Type"
    hdr_cells[5].text = "Team Response"

    # Loop through pages and extract annotations
    for page_num, page in enumerate(doc, start=1):
        for annot in page.annots() or []:
            author = annot.info.get("title", "Unknown")
            comment_text = annot.info.get("content", "")
            scope = annot.info.get("Subject", "")

            # Fit comments to the table
            row_cells = table.add_row().cells
            row_cells[0].text = author
            row_cells[1].text = str(page_num)
            row_cells[2].text = "N/A"  # PDF doesn't support paragraph numbers directly
            row_cells[3].text = scope if scope else "N/A"
            row_cells[4].text = comment_text
            row_cells[5].text = ""  # Empty response field

    # Save the document
    word_doc.save(word_output_path)
    print(f"Comments extracted and saved to {word_output_path}")


extract_comments_to_word("C:\Users\MUKANDI\Downloads\SAMM.pdf", "C:\Users\MUKANDI\OneDrive - UNHCR\Desktop\New folder (2)\extracted_comments.docx")
